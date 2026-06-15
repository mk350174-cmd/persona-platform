#!/usr/bin/env python3
"""
Parse Turkish HPEP-100 questions from Word document (.docx).

Input: Word document with 50 Turkish questions (expected June 15, 06:00 UTC)
Output: JSON with Turkish text for all 50 questions (S1-S50)

Usage:
    python scripts/parse_turkish_docx.py <input.docx> <output.json>

Handles:
    - UTF-8 encoding
    - Tables and formatted text
    - Question numbering (S1, S2, ..., S50)
    - Metadata tracking (file hash, parse time)
"""

import json
import sys
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Optional

def parse_docx_to_questions(docx_path: str) -> dict:
    """
    Parse a .docx file and extract 50 Turkish HPEP-100 questions.

    Returns a dict: {
        "metadata": {
            "file_hash": str,
            "file_size": int,
            "parse_time": str (ISO 8601),
            "total_questions": int,
        },
        "questions": {
            "S1": "Turkish question text",
            "S2": "Turkish question text",
            ...
            "S50": "Turkish question text",
        }
    }
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    if path.suffix.lower() != '.docx':
        raise ValueError(f"Expected .docx file, got: {path.suffix}")

    # Calculate file hash for integrity tracking
    file_hash = _calculate_hash(path)
    file_size = path.stat().st_size
    parse_time = datetime.utcnow().isoformat() + "Z"

    # Try to import python-docx; gracefully degrade if unavailable
    try:
        from docx import Document
        doc = Document(docx_path)
        questions = _extract_from_docx(doc)
    except ImportError:
        raise ImportError(
            "python-docx not installed. Install with: pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}")

    if len(questions) != 50:
        raise ValueError(
            f"Expected 50 questions, found {len(questions)}. "
            "Check file format or recount."
        )

    return {
        "metadata": {
            "file_hash": file_hash,
            "file_size": file_size,
            "parse_time": parse_time,
            "total_questions": len(questions),
        },
        "questions": questions,
    }


def _calculate_hash(path: Path) -> str:
    """Calculate SHA256 hash of file for integrity tracking."""
    sha256 = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def _extract_from_docx(doc) -> dict[str, str]:
    """
    Extract questions from a python-docx Document object.

    Searches for patterns:
    - "S1", "S2", ..., "S50" as question IDs
    - Treats paragraph or table cell following the ID as the question text
    """
    questions = {}
    all_text = []

    # Collect all text from paragraphs and tables
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            all_text.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text.strip())

    # Parse for question IDs (S1-S50)
    i = 0
    while i < len(all_text):
        line = all_text[i].strip()

        # Check if this line contains a question ID (S1, S2, ..., S50)
        qid = _extract_question_id(line)
        if qid:
            # Try to get the question text from the same line or the next line
            question_text = line

            # If the ID is at the start and line is just "S1", get next line as text
            if line.startswith(qid) and len(line) == len(qid):
                if i + 1 < len(all_text):
                    question_text = all_text[i + 1].strip()
                    i += 1
            else:
                # ID is embedded; extract text after it
                question_text = line[len(qid):].strip()
                if not question_text and i + 1 < len(all_text):
                    question_text = all_text[i + 1].strip()
                    i += 1

            if question_text:
                questions[qid] = question_text

        i += 1

    return questions


def _extract_question_id(text: str) -> Optional[str]:
    """
    Extract question ID (S1-S50) from text.
    Returns the ID if found (e.g., "S1"), else None.
    """
    text = text.strip()

    # Check prefix: "S1", "S1.", "S1:", "S1 ", etc.
    for i in range(1, 51):
        qid = f"S{i}"
        if text.startswith(qid) and (
            len(text) == len(qid) or
            text[len(qid)] in " .:(\n\t"
        ):
            return qid

    return None


def save_json(data: dict, output_path: str) -> None:
    """Save parsed questions to JSON with pretty formatting."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scripts/parse_turkish_docx.py <input.docx> <output.json>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    try:
        print(f"Parsing Turkish HPEP-100 from: {input_file}")
        result = parse_docx_to_questions(input_file)

        save_json(result, output_file)

        print(f"✓ Successfully parsed {result['metadata']['total_questions']} questions")
        print(f"✓ File hash: {result['metadata']['file_hash'][:16]}...")
        print(f"✓ Output saved to: {output_file}")

    except Exception as e:
        print(f"✗ Error: {e}", file=sys.stderr)
        sys.exit(1)
