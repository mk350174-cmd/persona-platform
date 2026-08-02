#!/usr/bin/env python3
"""
Parse KOMBİNASYON docx and extract all 250 hybrid persona combinations.

Output: data/hybrid_personas_raw.jsonl (one JSON per line)

Bug fix (T1-158, 2026-07-31): most paragraphs in the source docx are
prefixed with an invisible zero-width space (U+200B), left over from how
the text was pasted into Word. ``str.strip()`` does not remove it, so
``para_text.startswith("KOMBİNASYON")`` silently failed for 202 of the 250
combination headers — those combinations' text got swallowed into whatever
combo's "characteristic" field the parser was still filling in (e.g.
komb_001 originally absorbed the text of combos 2-5). Every text
comparison/extraction below now goes through ``_clean()``, which strips
zero-width space (and a stray BOM/no-break-space) before the normal
``.strip()``. Re-running against the same docx now finds all 250 headers
with no gaps (verified: numbers 1..250 present, none missing).
"""

import json
import re
from pathlib import Path
from typing import Dict, List, Any
from docx import Document

_INVISIBLE_CHARS = "​﻿\xa0"


def _clean(text: str) -> str:
    return text.strip().translate({ord(c): None for c in _INVISIBLE_CHARS}).strip()


def parse_active_layers(text: str) -> List[int]:
    """Extract K-layer numbers from 'Katman X' references in a paragraph."""
    pattern = r'Katman\s+(\d+)'
    matches = re.findall(pattern, text)

    layers = []
    for match in matches:
        layer_num = int(match)
        if 1 <= layer_num <= 98 and layer_num not in layers:
            layers.append(layer_num)

    return sorted(list(set(layers)))


def extract_all_personas(docx_path: str) -> List[Dict[str, Any]]:
    """Extract all persona combinations from DOCX file."""
    doc = Document(docx_path)

    personas = []
    i = 0
    while i < len(doc.paragraphs):
        para_text = _clean(doc.paragraphs[i].text)

        if para_text.startswith("KOMBİNASYON"):
            match = re.match(r'KOMBİNASYON\s+(\d+):\s*(.+?)\s*\((.+?)\)', para_text)
            if not match:
                i += 1
                continue

            combo_num = int(match.group(1))
            name_tr = match.group(2).strip()
            name_en = match.group(3).strip()

            combo = {
                "id": f"komb_{combo_num:03d}_{name_en.lower().replace(' ', '_').replace('-', '_')}",
                "number": combo_num,
                "name_tr": name_tr,
                "name_en": name_en,
                "use_case": "",
                "active_layers": [],
                "suppressed_layers": [],
                "characteristic": "",
            }

            i += 1
            while i < len(doc.paragraphs):
                para = _clean(doc.paragraphs[i].text)

                if para.startswith("KOMBİNASYON"):
                    break

                if "Kullanım Amacı:" in para:
                    combo["use_case"] = para.replace("Kullanım Amacı:", "").strip()
                elif "Aktif (Baskın) Katmanlar:" in para:
                    i += 1
                    while i < len(doc.paragraphs):
                        active_para = _clean(doc.paragraphs[i].text)
                        if active_para.startswith("KOMBİNASYON"):
                            break
                        if "Baskılanan Katmanlar:" in active_para:
                            combo["suppressed_layers"].extend(parse_active_layers(active_para))
                            break
                        if "Profil Karakteristiği:" in active_para:
                            break
                        if active_para:
                            combo["active_layers"].extend(parse_active_layers(active_para))
                        i += 1
                    continue
                elif "Baskılanan Katmanlar:" in para:
                    combo["suppressed_layers"].extend(parse_active_layers(para))
                elif "Profil Karakteristiği:" in para:
                    combo["characteristic"] = para.replace("Profil Karakteristiği:", "").strip()
                    i += 1
                    while i < len(doc.paragraphs):
                        char_para = _clean(doc.paragraphs[i].text)
                        if char_para.startswith("KOMBİNASYON"):
                            break
                        if char_para:
                            combo["characteristic"] += " " + char_para
                        i += 1
                    i -= 1

                i += 1

            combo["active_layers"] = sorted(list(set(combo["active_layers"])))
            combo["suppressed_layers"] = sorted(list(set(combo["suppressed_layers"])))

            if combo["name_tr"] and combo["number"]:
                personas.append(combo)
        else:
            i += 1

    return sorted(personas, key=lambda x: x["number"])


def main():
    docx_path = (
        Path(__file__).resolve().parent.parent
        / "docs" / "genesis" / "KOMBİNASYON_1-250_Hibrit_Persona_Katalogu.docx"
    )
    output_dir = Path(__file__).resolve().parent.parent / "data"
    output_dir.mkdir(parents=True, exist_ok=True)

    personas = extract_all_personas(str(docx_path))

    # Write to JSONL
    output_file = output_dir / "hybrid_personas_raw.jsonl"
    with open(output_file, 'w', encoding='utf-8') as f:
        for p in personas:
            f.write(json.dumps(p, ensure_ascii=False) + '\n')

    numbers = sorted(p["number"] for p in personas)
    missing = sorted(set(range(1, 251)) - set(numbers)) if numbers else list(range(1, 251))
    print(f"Extracted {len(personas)} personas")
    if missing:
        print(f"WARNING: missing combination numbers: {missing}")


if __name__ == "__main__":
    main()
